"""
文档生成器服务
生成符合格式要求的Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional
from pathlib import Path

from core.reference_formatter import ReferenceFormatter
from core.translator import translator
from services.logger import logger


class DocumentGenerator:
    """Word文档生成器"""

    def __init__(self):
        self.doc = Document()
        self.ref_formatter = ReferenceFormatter("vancouver")
        self.progress_callback = None

    def set_progress_callback(self, callback):
        """设置进度回调函数"""
        self.progress_callback = callback

    def _report_progress(self, message, progress=0):
        """报告进度"""
        if self.progress_callback:
            self.progress_callback(message, progress)

    def set_chinese_style(self):
        """设置中文样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Inches(0.3)

    def set_english_style(self):
        """设置英文样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Inches(0.3)

    def add_title(self, title: str, subtitle: str = ""):
        """添加标题"""
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(16)
            run.font.bold = True

        if subtitle:
            sub_heading = self.doc.add_paragraph(subtitle)
            sub_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in sub_heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    def add_section_heading(self, text: str, level: int = 1):
        """添加章节标题"""
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(14 if level == 1 else 12)

    def add_paragraph(self, text: str, indent: bool = True):
        """添加段落"""
        para = self.doc.add_paragraph()
        if indent:
            para.paragraph_format.first_line_indent = Inches(0.3)
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        return para

    def add_bilingual_paragraph(self, chinese: str, english: str):
        """添加中英文对照段落"""
        para = self.doc.add_paragraph()
        para.paragraph_format.first_line_indent = Inches(0.3)

        run_zh = para.add_run(chinese)
        run_zh.font.name = '宋体'
        run_zh._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_zh.font.size = Pt(12)

        para.add_run("\n")

        run_en = para.add_run(english)
        run_en.font.name = 'Times New Roman'
        run_en.font.size = Pt(12)

    def add_table(self, data: list, headers: list = None):
        """添加表格"""
        if not data:
            return

        rows = len(data) + (1 if headers else 0)
        cols = len(data[0]) if isinstance(data[0], list) else 1

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        current_row = 0

        if headers:
            for i, header in enumerate(headers):
                cell = table.rows[current_row].cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
            current_row += 1

        for row_data in data:
            if isinstance(row_data, list):
                for i, cell_data in enumerate(row_data):
                    cell = table.rows[current_row].cells[i]
                    cell.text = str(cell_data)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10)
            current_row += 1

        return table

    def add_references(self, references: list, style: str = "vancouver", language: str = "中文"):
        """添加参考文献"""
        self.doc.add_page_break()
        if language == "英文":
            self.add_section_heading("References", level=1)
        elif language == "中英文双语":
            self.add_section_heading("参考文献 / References", level=1)
        else:
            self.add_section_heading("参考文献", level=1)

        self.ref_formatter = ReferenceFormatter(style)

        for i, ref in enumerate(references, 1):
            formatted = self.ref_formatter.format_reference(ref)
            para = self.doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(-0.3)
            para.paragraph_format.left_indent = Inches(0.3)

            run = para.add_run(f"{i}. {formatted}")
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)

    def save(self, output_path: str):
        """保存文档"""
        self.doc.save(output_path)

    def generate_paper(
        self,
        session,
        output_path: str,
        language: str = "中文",
        ref_style: str = "vancouver"
    ):
        """生成完整论文"""

        # 标准化语言参数，兼容前端传递的"English"和"英文"
        if language.lower() == "english" or language == "英文":
            language = "英文"
        elif language.lower() == "chinese" or language == "中文":
            language = "中文"

        self._report_progress("开始生成文档...", 5)

        if language == "中英文双语":
            self.set_chinese_style()
        elif language == "中文":
            self.set_chinese_style()
        else:
            self.set_english_style()

        title = session.research_topic or session.research_question or "医学论文"
        research_context = title
        if language == "英文":
            self._report_progress("正在翻译标题...", 10)
            title = translator.translate_to_english(title, research_context)
        self.add_title(title, "")

        self.doc.add_paragraph()
        self.add_paragraph("[作者姓名]", indent=False)
        self.add_paragraph("[单位地址]", indent=False)
        self.doc.add_paragraph()

        if language == "英文":
            self.add_section_heading("Abstract", level=1)
            self.add_paragraph("[Abstract content to be generated]")
            self.doc.add_paragraph()
            self.add_section_heading("Keywords", level=2)
            self.add_paragraph("[Keyword 1, Keyword 2, Keyword 3]", indent=False)
        elif language == "中英文双语":
            self.add_section_heading("摘要 / Abstract", level=1)
            self.add_paragraph("[摘要内容]")
            self.add_paragraph("[Abstract content]")
            self.doc.add_paragraph()
            self.add_section_heading("关键词 / Keywords", level=2)
            self.add_paragraph("[关键词1, 关键词2, 关键词3]", indent=False)
        else:
            self.add_section_heading("摘要", level=1)
            self.add_paragraph("[摘要内容]")
            self.doc.add_paragraph()
            self.add_section_heading("关键词", level=2)
            self.add_paragraph("[关键词1, 关键词2, 关键词3]", indent=False)
        self.doc.add_paragraph()
        
        self._report_progress("正在处理章节内容...", 15)

        if session.outline:
            self._generate_content_from_outline(session.outline, language, research_context)
        else:
            self._generate_sections(session.sections, language, research_context)

        self._report_progress("正在添加参考文献...", 90)
        self.add_references([], style=ref_style, language=language)

        self._report_progress("正在保存文档...", 95)
        self.save(output_path)
        
        self._report_progress("文档生成完成", 100)

    def _generate_content_from_outline(self, outline_node, language="中文", research_context=""):
        """从提纲生成内容"""
        for child in outline_node.children:
            skip_title = False
            
            if child.title and (
                "核心内容" in child.title or 
                "关键要点" in child.title or
                child.title.strip().endswith("核心内容") or
                child.title.strip().endswith("关键要点")
            ):
                skip_title = True
            
            if not skip_title:
                # 根据语言选择标题
                display_title = child.title
                if language == "英文":
                    display_title = child.title_en or translator.translate_outline_title(child.title, research_context)
                
                if child.level == 1:
                    self.add_section_heading(display_title, level=1)
                elif child.level == 2:
                    self.add_section_heading(display_title, level=2)
                elif child.level == 3:
                    self.add_section_heading(display_title, level=3)

            # 根据语言选择内容，必要时翻译
            display_content = ""
            if language == "中英文双语" and child.content_zh and child.content_en:
                self.add_bilingual_paragraph(child.content_zh, child.content_en)
            elif language == "英文":
                logger.debug(f"[document_generator] translating to English", section_title=child.title, has_content_en=bool(child.content_en), has_content_zh=bool(child.content_zh))
                if child.content_en:
                    display_content = child.content_en
                    logger.debug(f"[document_generator] using existing English content", content_length=len(display_content))
                elif child.content_zh:
                    # 翻译中文内容
                    logger.info(f"[document_generator] translating Chinese content to English", content_length=len(child.content_zh))
                    display_content = translator.translate_to_english(child.content_zh, research_context)
                    logger.info(f"[document_generator] translation completed", translated_length=len(display_content))
                if display_content:
                    self.add_paragraph(display_content)
            else:
                # 中文模式
                if child.content_zh:
                    display_content = child.content_zh
                elif child.content_en:
                    display_content = translator.translate_to_chinese(child.content_en, research_context)
                if display_content:
                    self.add_paragraph(display_content)

            if child.children:
                self._generate_content_from_outline(child, language, research_context)

    def _generate_sections(self, sections: list, language="中文", research_context=""):
        """从章节列表生成内容"""
        logger.debug(f"[document_generator] _generate_sections called, language={language}, sections_count={len(sections)}")
        
        total_sections = len(sections)
        for i, section in enumerate(sections):
            if isinstance(section, dict):
                # 支持多种字段名
                title = section.get("title") or section.get("chapter") or "章节"
                content_zh = section.get("content_zh", "")
                content_en = section.get("content_en", "")
                content = section.get("content", "")
            else:
                title = "章节"
                content_zh = ""
                content_en = ""
                content = str(section)

            # 计算当前进度 (15% - 85%)
            progress = 15 + int((i / total_sections) * 70)
            self._report_progress(f"正在处理章节: {title} ({i+1}/{total_sections})", progress)

            logger.debug(f"[document_generator] Processing section {i}: title={title[:30]}..., content_zh_len={len(content_zh)}, content_en_len={len(content_en)}, content_len={len(content)}")

            # 根据语言选择标题
            display_title = title
            if language == "英文":
                logger.debug(f"[document_generator] Translating title to English: {title[:30]}...")
                display_title = translator.translate_outline_title(title, research_context)
                logger.debug(f"[document_generator] Translated title: {display_title[:30]}...")

            # 根据语言选择内容
            display_content = ""
            if language == "中英文双语" and content_zh and content_en:
                self.add_section_heading(display_title, level=2)
                self.add_bilingual_paragraph(content_zh, content_en)
                continue
            elif language == "英文":
                self._report_progress(f"正在翻译章节内容: {title}", progress + 2)
                logger.debug(f"[document_generator] English mode: selecting content")
                if content_en:
                    logger.debug(f"[document_generator] Using existing content_en")
                    display_content = content_en
                elif content_zh:
                    logger.debug(f"[document_generator] Translating content_zh to English (len={len(content_zh)})")
                    display_content = translator.translate_to_english(content_zh, research_context)
                    logger.debug(f"[document_generator] Translated content length: {len(display_content)}")
                elif content:
                    logger.debug(f"[document_generator] Translating content to English (len={len(content)})")
                    display_content = translator.translate_to_english(content, research_context)
                    logger.debug(f"[document_generator] Translated content length: {len(display_content)}")
            else:
                # 中文模式
                logger.debug(f"[document_generator] Chinese mode: selecting content")
                if content_zh:
                    display_content = content_zh
                elif content:
                    display_content = content
                elif content_en:
                    display_content = translator.translate_to_chinese(content_en, research_context)

            self.add_section_heading(display_title, level=2)
            if display_content:
                self.add_paragraph(display_content)
