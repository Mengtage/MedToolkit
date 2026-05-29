
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os

class WordGenerator:
    def __init__(self, output_path: str):
        self.doc = Document()
        self.output_path = output_path
        self._set_default_style()
    
    def _set_default_style(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Inches(0.3)
    
    def _add_chinese_text(self, text: str, paragraph=None):
        if paragraph is None:
            paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        return paragraph
    
    def _add_english_text(self, text: str, paragraph=None):
        if paragraph is None:
            paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return paragraph
    
    def add_bilingual_paragraph(self, chinese_text: str, english_text: str):
        para = self.doc.add_paragraph()
        self._add_chinese_text(chinese_text, para)
        self._add_english_text(f"\n{english_text}", para)
    
    def add_heading(self, chinese_text: str, english_text: str, level: int = 1):
        heading = self.doc.add_heading(level=level)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._add_chinese_text(chinese_text, heading)
        self._add_english_text(f" / {english_text}", heading)
    
    def add_title(self, chinese_title: str, english_title: str):
        title = self.doc.add_heading(level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._add_chinese_text(chinese_title, title)
        self._add_english_text(f"\n{english_title}", title)
    
    def add_authors(self, authors: str):
        para = self.doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._add_chinese_text(authors, para)
    
    def add_abstract(self, chinese_abstract: str, english_abstract: str):
        self.add_heading("摘要", "Abstract", level=1)
        self.add_bilingual_paragraph(chinese_abstract, english_abstract)
    
    def add_keywords(self, chinese_keywords: str, english_keywords: str):
        self.add_heading("关键词", "Keywords", level=2)
        self.add_bilingual_paragraph(chinese_keywords, english_keywords)
    
    def add_section(self, chinese_title: str, english_title: str, chinese_content: str, english_content: str):
        self.add_heading(chinese_title, english_title, level=1)
        self.add_bilingual_paragraph(chinese_content, english_content)
    
    def add_table(self, title: str, data):
        self.doc.add_heading(title, level=2)
        if data:
            table = self.doc.add_table(rows=len(data), cols=len(data[0]))
            for i, row in enumerate(data):
                for j, cell in enumerate(row):
                    table.rows[i].cells[j].text = str(cell)
    
    def add_figure_legend(self, figure_number: int, chinese_title: str, english_title: str, chinese_description: str, english_description: str):
        self.doc.add_heading(f"图{figure_number}", f"Figure {figure_number}", level=2)
        self.add_bilingual_paragraph(chinese_title, english_title)
        self.add_bilingual_paragraph(chinese_description, english_description)
    
    def add_supplement_table(self, table_number: int, chinese_title: str, english_title: str, chinese_description: str, english_description: str):
        self.doc.add_heading(f"附表{table_number}", f"Supplement Table {table_number}", level=2)
        self.add_bilingual_paragraph(chinese_title, english_title)
        self.add_bilingual_paragraph(chinese_description, english_description)
    
    def save(self):
        self.doc.save(self.output_path)

