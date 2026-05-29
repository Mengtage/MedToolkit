"""
文件解析模块
支持从多种文件格式提取文本内容
"""

import io
import os
import subprocess
import tempfile


def extract_text_from_file(content: bytes, filename: str) -> str:
    """
    从各种文件格式中提取文本内容

    Args:
        content: 文件二进制内容
        filename: 文件名（用于判断扩展名）

    Returns:
        提取的文本内容

    Raises:
        ValueError: 不支持的文件类型或解析失败
    """
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    if ext in ('txt', 'md'):
        return content.decode('utf-8', errors='replace')

    elif ext == 'csv':
        try:
            import csv
            text_parts = []
            reader = csv.reader(io.StringIO(content.decode('utf-8', errors='replace')))
            for row in reader:
                row_str = " | ".join([cell.strip() for cell in row])
                if row_str.strip():
                    text_parts.append(row_str)
            return "\n".join(text_parts) if text_parts else "（CSV文件内容提取为空）"
        except Exception as e:
            raise ValueError(f"CSV解析失败: {e}")

    elif ext == 'pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text.strip() or "（PDF文件内容提取为空）"
        except ImportError:
            raise ValueError("PDF解析需要安装PyMuPDF: pip install PyMuPDF")
        except Exception as e:
            raise ValueError(f"PDF解析失败: {e}")

    elif ext == 'docx':
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            
            # 提取段落文本
            text_parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)
            
            # 提取表格内容
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        # 合并单元格可能导致重复，需要去重
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    if row_cells:
                        table_text.append(" | ".join(row_cells))
                if table_text:
                    text_parts.append("【表格】")
                    text_parts.extend(table_text)
            
            full_text = "\n".join(text_parts)
            return full_text.strip() or "（Word文件内容提取为空）"
        except ImportError:
            raise ValueError("Word解析需要安装python-docx: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Word解析失败: {e}")

    elif ext == 'doc':
        try:
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            try:
                result = subprocess.run(
                    ['textutil', '-convert', 'txt', '-stdout', tmp_path],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout:
                    return result.stdout.strip()
                raise ValueError("textutil转换失败")
            finally:
                os.unlink(tmp_path)
        except FileNotFoundError:
            raise ValueError("macOS 'textutil'命令不可用，请将文档保存为.docx格式")
        except subprocess.TimeoutExpired:
            raise ValueError("文档转换超时，请将文档保存为.docx格式")
        except Exception as e:
            raise ValueError(f"旧版Word文档(.doc)解析失败: {e}")

    elif ext in ('xlsx', 'xls'):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            text_parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows_text = []
                for row in ws.iter_rows(values_only=True):
                    row_str = " | ".join([str(cell) for cell in row if cell is not None])
                    if row_str.strip():
                        rows_text.append(row_str)
                if rows_text:
                    text_parts.append(f"【工作表: {sheet_name}】")
                    text_parts.extend(rows_text)
            wb.close()
            return "\n".join(text_parts) if text_parts else "（Excel文件内容提取为空）"
        except ImportError:
            raise ValueError("Excel解析需要安装openpyxl: pip install openpyxl")
        except Exception as e:
            raise ValueError(f"Excel解析失败: {e}")

    else:
        raise ValueError(f"不支持的文件类型: .{ext}，支持的类型: .txt, .md, .csv, .docx, .doc, .pdf, .xlsx, .xls")