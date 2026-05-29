"""
文档生成器服务
生成符合格式要求的Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional
from pathlib import Path

from services.reference_formatter import ReferenceFormatter


class DocumentGenerator:
    """Word文档生成器"""

    def __init__(self):
        self.doc = Document()
        self.ref_formatter = ReferenceFormatter("vancouver")

    def set_chinese_style(self):
        """设置中文样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Inches(0.3)

    def set_english_style(self):
        """设置英文样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Inches(0.3)

    def add_title(self, title: str, subtitle: str = ""):
        """添加标题"""
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(16)
            run.font.bold = True

        if subtitle:
            sub_heading = self.doc.add_paragraph(subtitle)
            sub_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in sub_heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    def add_section_heading(self, text: str, level: int = 1):
        """添加章节标题"""
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(14 if level == 1 else 12)

    def add_paragraph(self, text: str, indent: bool = True):
        """添加段落"""
        para = self.doc.add_paragraph()
        if indent:
            para.paragraph_format.first_line_indent = Inches(0.3)
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        return para

    def add_bilingual_paragraph(self, chinese: str, english: str):
        """添加中英文对照段落"""
        para = self.doc.add_paragraph()
        para.paragraph_format.first_line_indent = Inches(0.3)

        run_zh = para.add_run(chinese)
        run_zh.font.name = '宋体'
        run_zh._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_zh.font.size = Pt(12)

        para.add_run("\n")

        run_en = para.add_run(english)
        run_en.font.name = 'Times New Roman'
        run_en.font.size = Pt(12)

    def add_table(self, data: list, headers: list = None):
        """添加表格"""
        if not data:
            return

        rows = len(data) + (1 if headers else 0)
        cols = len(data[0]) if isinstance(data[0], list) else 1

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        current_row = 0

        if headers:
            for i, header in enumerate(headers):
                cell = table.rows[current_row].cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
            current_row += 1

        for row_data in data:
            if isinstance(row_data, list):
                for i, cell_data in enumerate(row_data):
                    cell = table.rows[current_row].cells[i]
                    cell.text = str(cell_data)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10)
            current_row += 1

        return table

    def add_references(self, references: list, style: str = "vancouver"):
        """添加参考文献"""
        self.doc.add_page_break()
        self.add_section_heading("参考文献 / References", level=1)

        self.ref_formatter = ReferenceFormatter(style)

        for i, ref in enumerate(references, 1):
            formatted = self.ref_formatter.format_reference(ref)
            para = self.doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(-0.3)
            para.paragraph_format.left_indent = Inches(0.3)

            run = para.add_run(f"{i}. {formatted}")
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)

    def save(self, output_path: str):
        """保存文档"""
        self.doc.save(output_path)

    def generate_paper(
        self,
        session,
        output_path: str,
        language: str = "中文",
        ref_style: str = "vancouver"
    ):
        """生成完整论文"""

        if language == "中英文双语":
            self.set_chinese_style()
        elif language == "中文":
            self.set_chinese_style()
        else:
            self.set_english_style()

        title = session.research_topic or session.research_question or "医学论文"
        self.add_title(title, "")

        self.doc.add_paragraph()
        self.add_paragraph("[作者姓名]", indent=False)
        self.add_paragraph("[单位地址]", indent=False)
        self.doc.add_paragraph()

        self.add_section_heading("摘要 / Abstract", level=1)
        self.add_paragraph("[此处生成摘要内容]")
        self.doc.add_paragraph()

        self.add_section_heading("关键词 / Keywords", level=2)
        self.add_paragraph("[关键词1, 关键词2, 关键词3]", indent=False)
        self.doc.add_paragraph()

        if session.outline:
            self._generate_content_from_outline(session.outline, language)
        else:
            self._generate_sections(session.sections)

        self.add_references([], style=ref_style)

        self.save(output_path)

    def _generate_content_from_outline(self, outline_node, language="中文"):
        """从提纲生成内容"""
        for child in outline_node.children:
            skip_title = False
            
            if child.title and (
                "核心内容" in child.title or 
                "关键要点" in child.title or
                child.title.strip().endswith("核心内容") or
                child.title.strip().endswith("关键要点")
            ):
                skip_title = True
            
            if not skip_title:
                if child.level == 1:
                    self.add_section_heading(child.title, level=1)
                elif child.level == 2:
                    self.add_section_heading(child.title, level=2)
                elif child.level == 3:
                    self.add_section_heading(child.title, level=3)

            if child.content_zh:
                if language == "中英文双语" and child.content_en:
                    self.add_bilingual_paragraph(child.content_zh, child.content_en)
                else:
                    self.add_paragraph(child.content_zh)

            if child.children:
                self._generate_content_from_outline(child, language)

    def _generate_sections(self, sections: list):
        """从章节列表生成内容"""
        for section in sections:
            if isinstance(section, dict):
                title = section.get("title", "章节")
                content = section.get("content_zh", section.get("content", ""))
            else:
                title = "章节"
                content = str(section)

            self.add_section_heading(title, level=2)
            self.add_paragraph(content)
